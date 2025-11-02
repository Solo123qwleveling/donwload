import sys
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QTextEdit, QPushButton, QLabel,
                             QLineEdit, QSpinBox, QColorDialog, QGroupBox,
                             QGridLayout, QMessageBox, QFileDialog, QScrollArea)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont, QColor, QPalette, QSyntaxHighlighter, QTextCharFormat
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class JavaHighlighter(QSyntaxHighlighter):
    """Syntax highlighter for Java code preview"""

    def __init__(self, parent, colors):
        super().__init__(parent)
        self.colors = colors

    def highlightBlock(self, text):
        format = QTextCharFormat()

        trimmed = text.strip()
        if trimmed.startswith("//"):
            format.setForeground(QColor(self.colors['comment']))
        elif trimmed.startswith("@"):
            format.setForeground(QColor(self.colors['annotation']))
        elif "case " in text:
            format.setForeground(QColor("#FF9696"))
        elif any(kw in text for kw in ["public ", "class ", "private ", "protected "]):
            format.setForeground(QColor(self.colors['keyword']))
        elif '"' in text:
            format.setForeground(QColor(self.colors['string']))
        else:
            format.setForeground(QColor(self.colors['default']))

        self.setFormat(0, len(text), format)


class ColorButton(QPushButton):
    """Custom color picker button"""

    def __init__(self, color, parent=None):
        super().__init__(parent)
        self.color = QColor(color)
        self.setFixedHeight(40)
        self.update_color()
        self.clicked.connect(self.choose_color)

    def choose_color(self):
        color = QColorDialog.getColor(self.color, self, "Choose Color")
        if color.isValid():
            self.color = color
            self.update_color()

    def update_color(self):
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.color.name()};
                border: 2px solid rgba(147, 51, 234, 0.3);
                border-radius: 8px;
                color: {'#000' if self.color.lightness() > 128 else '#fff'};
                font-weight: bold;
                padding: 8px;
            }}
            QPushButton:hover {{
                border: 2px solid rgba(147, 51, 234, 0.6);
            }}
        """)
        self.setText(self.color.name().upper())

    def get_color(self):
        return self.color.name()


class JavaCodeFormatter(QMainWindow):
    """Main application window for Java Code to Word Formatter"""

    def __init__(self):
        super().__init__()
        self.colors = {
            'background': '#000000',
            'comment': '#00FF00',
            'annotation': '#FFB450',
            'keyword': '#FF6464',
            'string': '#50C8FF',
            'default': '#FFFFFF'
        }
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Java Code to Word Formatter")
        self.setGeometry(100, 100, 1400, 900)

        # Set modern dark theme with gradient-like effect
        self.setStyleSheet("""
            QMainWindow {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #1e1b4b, stop:0.5 #581c87, stop:1 #1e1b4b);
            }
            QTextEdit {
                background-color: rgba(30, 41, 59, 0.5);
                color: #ffffff;
                border: 1px solid rgba(147, 51, 234, 0.3);
                border-radius: 12px;
                padding: 8px;
                font-family: Consolas, Monaco, monospace;
            }
            QTextEdit:focus {
                border: 2px solid rgba(147, 51, 234, 0.6);
            }
            QLineEdit {
                background-color: rgba(30, 41, 59, 0.5);
                color: #ffffff;
                border: 1px solid rgba(147, 51, 234, 0.3);
                border-radius: 8px;
                padding: 8px;
                font-size: 13px;
            }
            QLineEdit:focus {
                border: 2px solid rgba(147, 51, 234, 0.6);
            }
            QLabel {
                color: #e9d5ff;
                font-size: 13px;
            }
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #7c3aed, stop:1 #db2777);
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 20px;
                font-weight: bold;
                font-size: 13px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #6d28d9, stop:1 #be185d);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #5b21b6, stop:1 #9f1239);
            }
            QGroupBox {
                color: #e9d5ff;
                border: 1px solid rgba(147, 51, 234, 0.3);
                border-radius: 12px;
                margin-top: 15px;
                padding-top: 15px;
                font-weight: bold;
                background-color: rgba(255, 255, 255, 0.05);
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 8px;
                background-color: rgba(30, 41, 59, 0.8);
                border-radius: 4px;
            }
            QSpinBox {
                background-color: rgba(30, 41, 59, 0.5);
                color: #ffffff;
                border: 1px solid rgba(147, 51, 234, 0.3);
                border-radius: 8px;
                padding: 8px;
            }
        """)

        # Central widget with scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        central_widget = QWidget()
        scroll.setWidget(central_widget)
        self.setCentralWidget(scroll)

        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(20)

        # Header
        header = QWidget()
        header_layout = QVBoxLayout(header)
        header_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        title = QLabel("☕ Java Code to Word Formatter")
        title.setStyleSheet("font-size: 32px; font-weight: bold; color: white;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)

        subtitle = QLabel("Format your Java code with syntax highlighting for Word documents")
        subtitle.setStyleSheet("font-size: 14px; color: #e9d5ff;")
        subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)

        header_layout.addWidget(title)
        header_layout.addWidget(subtitle)
        main_layout.addWidget(header)

        # Content area (two columns)
        content_layout = QHBoxLayout()
        content_layout.setSpacing(20)

        # Left Panel
        left_panel = self.create_left_panel()
        content_layout.addWidget(left_panel, 1)

        # Right Panel
        right_panel = self.create_right_panel()
        content_layout.addWidget(right_panel, 1)

        main_layout.addLayout(content_layout)

        # Initial preview update
        self.update_preview()

    def create_left_panel(self):
        """Create the left input panel"""
        panel = QGroupBox("📝 Input Settings")
        panel.setStyleSheet(panel.styleSheet() + "QGroupBox { padding: 20px; }")
        layout = QVBoxLayout(panel)
        layout.setSpacing(15)

        # Document title
        title_label = QLabel("Document Title:")
        title_label.setStyleSheet("font-weight: bold;")
        self.title_input = QLineEdit("Accessibility Service Java Code")
        layout.addWidget(title_label)
        layout.addWidget(self.title_input)

        # Code input
        code_label = QLabel("Java Code:")
        code_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(code_label)

        self.code_input = QTextEdit()
        self.code_input.setFont(QFont("Consolas", 11))
        self.code_input.setMinimumHeight(300)
        self.code_input.setPlainText("""    @Override
    public void onInterrupt() {
        //whatever
    }""")
        layout.addWidget(self.code_input)

        # Color settings
        color_group = QGroupBox("🎨 Color Settings")
        color_layout = QGridLayout()
        color_layout.setSpacing(10)

        self.color_buttons = {}
        color_labels = {
            'background': 'Background',
            'comment': 'Comments',
            'annotation': 'Annotations',
            'keyword': 'Keywords',
            'string': 'Strings',
            'default': 'Default Text'
        }

        row = 0
        for key, label_text in color_labels.items():
            label = QLabel(label_text + ":")
            button = ColorButton(self.colors[key])
            self.color_buttons[key] = button
            color_layout.addWidget(label, row, 0)
            color_layout.addWidget(button, row, 1)
            row += 1

        color_group.setLayout(color_layout)
        layout.addWidget(color_group)

        # Font size
        font_layout = QHBoxLayout()
        font_label = QLabel("Font Size (pt):")
        font_label.setStyleSheet("font-weight: bold;")
        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(8, 16)
        self.font_size_spin.setValue(11)
        self.font_size_spin.setFixedWidth(80)
        font_layout.addWidget(font_label)
        font_layout.addWidget(self.font_size_spin)
        font_layout.addStretch()
        layout.addLayout(font_layout)

        return panel

    def create_right_panel(self):
        """Create the right preview panel"""
        panel = QGroupBox("👁 Preview")
        panel.setStyleSheet(panel.styleSheet() + "QGroupBox { padding: 20px; }")
        layout = QVBoxLayout(panel)
        layout.setSpacing(15)

        # Preview area
        self.preview = QTextEdit()
        self.preview.setFont(QFont("Consolas", 11))
        self.preview.setReadOnly(True)
        self.preview.setMinimumHeight(400)
        layout.addWidget(self.preview)

        # Buttons
        button_layout = QVBoxLayout()
        button_layout.setSpacing(10)

        update_preview_btn = QPushButton("🔄 Update Preview")
        update_preview_btn.clicked.connect(self.update_preview)
        update_preview_btn.setMinimumHeight(45)
        button_layout.addWidget(update_preview_btn)

        generate_btn = QPushButton("📥 Download Python Script")
        generate_btn.clicked.connect(self.download_script)
        generate_btn.setMinimumHeight(45)
        button_layout.addWidget(generate_btn)

        generate_doc_btn = QPushButton("📄 Generate Word Document")
        generate_doc_btn.clicked.connect(self.generate_document)
        generate_doc_btn.setMinimumHeight(45)
        generate_doc_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #15803d, stop:1 #166534);
                font-size: 14px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #16a34a, stop:1 #15803d);
            }
        """)
        button_layout.addWidget(generate_doc_btn)

        info_label = QLabel("💡 Tip: Download the Python script or generate directly")
        info_label.setStyleSheet("font-size: 11px; color: #c4b5fd; font-style: italic;")
        info_label.setWordWrap(True)
        info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        button_layout.addWidget(info_label)

        layout.addLayout(button_layout)

        return panel

    def update_preview(self):
        """Update the code preview with syntax highlighting"""
        code = self.code_input.toPlainText()
        self.preview.clear()

        # Update colors from buttons
        for key, button in self.color_buttons.items():
            self.colors[key] = button.get_color()

        # Set background color
        bg_color = QColor(self.colors['background'])
        palette = self.preview.palette()
        palette.setColor(QPalette.ColorRole.Base, bg_color)
        self.preview.setPalette(palette)

        # Add syntax highlighting
        self.highlighter = JavaHighlighter(self.preview.document(), self.colors)
        self.preview.setPlainText(code)

    def hex_to_rgb(self, hex_color):
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))

    def generate_python_script(self):
        """Generate the Python script content"""
        code = self.code_input.toPlainText()
        title = self.title_input.text()
        font_size = self.font_size_spin.value()

        # Update colors
        for key, button in self.color_buttons.items():
            self.colors[key] = button.get_color()

        comment_rgb = self.hex_to_rgb(self.colors['comment'])
        annotation_rgb = self.hex_to_rgb(self.colors['annotation'])
        keyword_rgb = self.hex_to_rgb(self.colors['keyword'])
        string_rgb = self.hex_to_rgb(self.colors['string'])
        default_rgb = self.hex_to_rgb(self.colors['default'])
        bg_hex = self.colors['background'].lstrip('#')

        script = f'''from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Create new document
doc = Document()
doc.add_heading("{title}", level=1)

# Create a 1x1 table to act as a textbox
table = doc.add_table(rows=1, cols=1)
table.autofit = False
table.columns[0].width = Inches(6.5)
cell = table.cell(0, 0)

# Apply background shading to the cell
cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd %s w:fill="{bg_hex}"/>' % nsdecls('w')))

# Add padding
cell._tc.get_or_add_tcPr().append(parse_xml(
    r'<w:tcMar %s>'
    r'  <w:left w:w="240" w:type="dxa"/>'
    r'  <w:right w:w="240" w:type="dxa"/>'
    r'  <w:top w:w="120" w:type="dxa"/>'
    r'  <w:bottom w:w="120" w:type="dxa"/>'
    r'</w:tcMar>' % nsdecls('w')))

# Code content
code_text = """{code}"""

# Insert colored runs
p = cell.paragraphs[0]
for line in code_text.split("\\n"):
    run = p.add_run(line + "\\n")
    run.font.name = "Consolas"
    run.font.size = Pt({font_size})

    # Simple syntax coloring
    if line.strip().startswith("//"):
        run.font.color.rgb = RGBColor{comment_rgb}
    elif line.strip().startswith("@"):
        run.font.color.rgb = RGBColor{annotation_rgb}
    elif "case " in line:
        run.font.color.rgb = RGBColor(255, 150, 150)
    elif any(kw in line for kw in ["public ", "class ", "private ", "protected "]):
        run.font.color.rgb = RGBColor{keyword_rgb}
    elif '"' in line:
        run.font.color.rgb = RGBColor{string_rgb}
    else:
        run.font.color.rgb = RGBColor{default_rgb}

# Save file
output_path = "./design_code_java.docx"
doc.save(output_path)
print(f"Document saved to {{output_path}}")
'''
        return script

    def download_script(self):
        """Download the Python script"""
        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Python Script",
                "java_code_formatter.py",
                "Python Files (*.py)"
            )

            if file_path:
                script = self.generate_python_script()
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(script)

                QMessageBox.information(
                    self,
                    "Success",
                    f"Python script saved successfully to:\n{file_path}\n\nRun this script to generate the Word document."
                )

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to save script:\n{str(e)}"
            )

    def generate_document(self):
        """Generate the Word document directly"""
        try:
            # Create new document
            doc = Document()
            doc.add_heading(self.title_input.text(), level=1)

            # Create table
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(6.5)
            cell = table.cell(0, 0)

            # Update colors
            for key, button in self.color_buttons.items():
                self.colors[key] = button.get_color()

            # Apply background color
            bg_hex = self.colors['background'].lstrip('#')
            cell._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd %s w:fill="%s"/>' % (nsdecls('w'), bg_hex))
            )

            # Add padding
            cell._tc.get_or_add_tcPr().append(parse_xml(
                r'<w:tcMar %s>'
                r'  <w:left w:w="240" w:type="dxa"/>'
                r'  <w:right w:w="240" w:type="dxa"/>'
                r'  <w:top w:w="120" w:type="dxa"/>'
                r'  <w:bottom w:w="120" w:type="dxa"/>'
                r'</w:tcMar>' % nsdecls('w')
            ))

            # Get code
            code_text = self.code_input.toPlainText()

            # Insert colored runs
            p = cell.paragraphs[0]
            for line in code_text.split("\n"):
                run = p.add_run(line + "\n")
                run.font.name = "Consolas"
                run.font.size = Pt(self.font_size_spin.value())

                # Apply syntax coloring
                trimmed = line.strip()
                if trimmed.startswith("//"):
                    rgb = self.hex_to_rgb(self.colors['comment'])
                elif trimmed.startswith("@"):
                    rgb = self.hex_to_rgb(self.colors['annotation'])
                elif "case " in line:
                    rgb = (255, 150, 150)
                elif any(kw in line for kw in ["public ", "class ", "private ", "protected "]):
                    rgb = self.hex_to_rgb(self.colors['keyword'])
                elif '"' in line:
                    rgb = self.hex_to_rgb(self.colors['string'])
                else:
                    rgb = self.hex_to_rgb(self.colors['default'])

                run.font.color.rgb = RGBColor(*rgb)

            # Save file
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Document",
                "design_code_java.docx",
                "Word Documents (*.docx)"
            )

            if file_path:
                doc.save(file_path)
                QMessageBox.information(
                    self,
                    "Success",
                    f"Document saved successfully to:\n{file_path}"
                )

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to generate document:\n{str(e)}"
            )


def main():
    app = QApplication(sys.argv)
    window = JavaCodeFormatter()
    window.show()
    sys.exit(app.exec())


if __name__ == '__main__':
    main()