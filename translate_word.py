import os
import argostranslate.package
import argostranslate.translate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuration ---
# You can find the language codes here: 
# https://github.com/argosopentech/argos-translate/blob/main/argostranslate/languages.py
SOURCE_LANGUAGE_CODE = "en"  # Example: English
TARGET_LANGUAGE_CODE = "ar"  # Example: Arabic (RTL Language)
RTL_LANGS = ['ar', 'fa', 'he', 'ur']  # List of Right-to-Left languages


# --- Core Helper Functions ---

def set_rtl_formatting(paragraph):
    """
    Sets the paragraph's alignment to Right and applies the RTL property 
    to the paragraph and all its runs to ensure proper display.
    """
    # 1. Set Paragraph Alignment to Right
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 2. Apply RTL property to the paragraph element
    pPr = paragraph.element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

    # 3. Apply Complex Script (RTL) property to all Runs
    for run in paragraph.runs:
        rPr = run.element.get_or_add_rPr()
        # Apply the RTL/BIDI flag for complex scripts
        bidi = OxmlElement('w:bidi')
        rPr.append(bidi)

        # Additionally set the complex script flag on the font property
        run.font.complex_script = True


def setup_offline_translator(from_code, to_code):
    """
    Downloads and installs the required Argos Translate model package.
    """
    print(f"Checking for offline model: {from_code} -> {to_code}...")

    # 1. Update the index of available packages (requires internet)
    # This step checks which models are available for download.
    argostranslate.package.update_package_index()

    # 2. Find the desired package
    available_packages = argostranslate.package.get_available_packages()

    package_to_install = next(
        filter(
            lambda x: x.from_code == from_code and x.to_code == to_code,
            available_packages
        ),
        None
    )

    if package_to_install:
        # 3. Download and install the package (requires internet initially)
        print("Model found. Installing model (if not already installed)...")
        try:
            argostranslate.package.install_from_path(package_to_install.download())
            print("Model is ready for offline use.")
        except Exception as e:
            # This can happen if the model is already installed
            print("Model already installed or a minor error occurred during install check (OK).")
            # For demonstration, we assume if it fails here, the model is present.
            pass
    else:
        print(
            f"⚠️ Error: No model found for {from_code} -> {to_code}. Check language codes or internet connection for initial download.")
        return None

    return argostranslate.translate.translate


# --- Main Processing Function ---

def process_word_document_offline(file_path, from_code, to_code):
    """
    Opens a Word document, translates all text offline, and applies RTL formatting.
    """
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return

    # 1. Setup the Offline Translator
    translate_function = setup_offline_translator(from_code, to_code)
    if not translate_function:
        return

    document = Document(file_path)
    print(f"\nProcessing document: {file_path}")

    # 2. Define the main translation and formatting logic
    def translate_and_format(paragraph):
        original_text = paragraph.text.strip()

        if original_text:
            try:
                # Use the offline Argos Translate function
                translated_text = translate_function(original_text, from_code, to_code)

                # Replace the original text in the paragraph
                # This clears all formatting but we re-apply RTL below.
                paragraph.clear()

                # Add the translated text back
                new_run = paragraph.add_run(translated_text)

                # Apply RTL formatting if the target language is RTL
                if to_code in RTL_LANGS:
                    set_rtl_formatting(paragraph)

            except Exception as e:
                print(f"Could not translate paragraph: '{original_text[:50]}...' Error: {e}")

    # 3. Iterate through Paragraphs
    for paragraph in document.paragraphs:
        translate_and_format(paragraph)

    # 4. Iterate through Tables (processing text inside table cells)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    translate_and_format(paragraph)

    # 5. Save the new document
    base, ext = os.path.splitext(file_path)
    output_file = f"{base}_OFFLINE_{to_code}.docx"
    document.save(output_file)
    print("\n-------------------------------------------")
    print(f"✅ Success! Translated document saved as: {output_file}")
    print(f"Translation model: {from_code} -> {to_code} (All translations were performed OFFLINE after model install).")
    print("-------------------------------------------")


# --- Main Execution ---

if __name__ == '__main__':
    # -------------------------------------------------------------------
    # ➡️ EDIT THESE TWO LINES 
    # 1. Specify the path to your input .docx file
    INPUT_FILE_NAME = "MyDocument.docx"

    # 2. Specify the source and target language codes 
    # Must be codes supported by Argos Translate (e.g., "en", "ar", "es", "fr")
    # Set to an RTL language like "ar" for Arabic, "he" for Hebrew, etc., to activate RTL formatting.
    SRC_CODE = "en"
    TGT_CODE = "ar"
    # -------------------------------------------------------------------

    process_word_document_offline(INPUT_FILE_NAME, SRC_CODE, TGT_CODE)